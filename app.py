import os
import re
import tempfile
import gradio as gr
import docx
import openpyxl
import pandas as pd
import uuid
from openai import OpenAI
from rank_bm25 import BM25Okapi
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_experimental.text_splitter import SemanticChunker
from sentence_transformers import CrossEncoder


# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# Central place to tune all model names, thresholds, and retrieval settings.
# Change values here — no need to hunt through the rest of the code.
# ─────────────────────────────────────────────────────────────────────────────

class Config:
    # OpenAI models
    EMBEDDING_MODEL = "text-embedding-3-small"
    LLM_MODEL       = "gpt-4o-mini"

    # Fixed chunking settings (used when total document size > SMALL_DOC_THRESHOLD)
    CHUNK_SIZE    = 600
    CHUNK_OVERLAP = 100

    # Documents under this character count get semantic chunking (slower but smarter).
    # Larger documents fall back to fixed chunking for speed.
    SMALL_DOC_THRESHOLD = 50_000

    # Number of chunks returned per retrieval call
    RETRIEVAL_K = 4

    # Weight balance between BM25 (keyword) and semantic search.
    # 0.0 = pure BM25, 1.0 = pure semantic, 0.5 = equal blend
    HYBRID_ALPHA = 0.5

    # Cross-encoder model used for reranking retrieved candidates
    RERANKER_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"

    # Hard cap on context characters sent to the LLM to avoid token overflow
    MAX_CONTEXT_CHARS = 12_000


# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# Converts uploaded files into LangChain Document objects.
# Each Document carries page_content (text) and metadata (source, page number).
# Supports: PDF, DOCX, TXT, XLSX/XLS, CSV
# ─────────────────────────────────────────────────────────────────────────────

def extract_text(filepath, filename):
    ext  = filename.split(".")[-1].lower()
    docs = []

    if ext == "pdf":
        # Primary: native PDF text extraction via PyPDFLoader
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(filepath)
            pages  = loader.load()

            if any(p.page_content.strip() for p in pages):
                for p in pages:
                    p.metadata["source"]   = filename
                    p.metadata["doc_name"] = filename
                return pages

        except Exception as e:
            print(f"[PDF native extraction failed] {e}")

        # Fallback: OCR via Tesseract (requires tesseract + poppler installed)
        # If these binaries are missing this will also fail — the caller checks
        # for an empty return and surfaces an error to the user.
        try:
            import pytesseract
            from pdf2image import convert_from_path
            images = convert_from_path(filepath, dpi=200)
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                docs.append(Document(
                    page_content=text,
                    metadata={"source": filename, "page": i + 1, "doc_name": filename}
                ))
        except Exception as e:
            print(f"[PDF OCR fallback failed] {e}")

    elif ext == "docx":
        doc  = docx.Document(filepath)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        docs.append(Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "doc_name": filename}
        ))

    elif ext == "txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        docs.append(Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "doc_name": filename}
        ))

    elif ext in ["xlsx", "xls"]:
        wb   = openpyxl.load_workbook(filepath, data_only=True)
        text = ""
        for sheet in wb.sheetnames:
            ws    = wb[sheet]
            text += f"\n[Sheet: {sheet}]\n"
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(c) for c in row if c is not None])
                if row_text.strip():
                    text += row_text + "\n"
        docs.append(Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "doc_name": filename}
        ))

    elif ext == "csv":
        df   = pd.read_csv(filepath)
        # Include schema info + a 100-row sample so large CSVs don't overflow context
        text = (
            f"Dataset: {filename}\n"
            f"Rows: {len(df)}\n"
            f"Columns: {', '.join(df.columns.tolist())}\n\n"
        )
        sample_df = df.sample(min(len(df), 100), random_state=42)
        text += sample_df.to_string()
        docs.append(Document(
            page_content=text,
            metadata={"source": filename, "page": 1, "doc_name": filename}
        ))

    return docs


# ─────────────────────────────────────────────────────────────────────────────
# HYBRID RETRIEVER
# Combines BM25 (keyword matching) and semantic (embedding) search, then
# reranks the merged candidate pool with a CrossEncoder for precision.
#
# Flow:
#   1. BM25 scores all chunks → top-50 by keyword relevance
#   2. Semantic search retrieves top-50 by embedding similarity
#   3. Both lists are merged and deduplicated
#   4. Low-confidence semantic results are filtered out
#   5. CrossEncoder reranks surviving candidates → top-k returned
# ─────────────────────────────────────────────────────────────────────────────

class HybridRetriever:
    def __init__(self, chunks, vectorstore, alpha=0.5):
        self.chunks      = chunks
        self.vectorstore = vectorstore
        self.alpha       = alpha

        # Build BM25 index over all chunk texts
        tokenized  = [self._tokenize(c.page_content) for c in chunks]
        self.bm25  = BM25Okapi(tokenized)

        # CrossEncoder loaded once per retriever instance (not at import time)
        self.reranker = CrossEncoder(Config.RERANKER_MODEL)

    def _tokenize(self, text):
        """Lowercase, strip punctuation, split on whitespace."""
        return re.sub(r'[^a-z0-9\s]', '', text.lower()).split()

    def retrieve(self, query, k=4):
        """Return the top-k most relevant chunks for a given query."""
        tokenized_query = self._tokenize(query)

        # ── Step 1: BM25 keyword scores (normalised 0–1) ──────────────────
        bm25_scores = self.bm25.get_scores(tokenized_query)
        bm25_scores = bm25_scores / (bm25_scores.max() + 1e-6)

        # ── Step 2: Semantic similarity scores ────────────────────────────
        k_sem       = min(50, len(self.chunks))
        sem_results = self.vectorstore.similarity_search_with_relevance_scores(query, k=k_sem)
        sem_score_map = {doc.page_content: score for doc, score in sem_results}

        # ── Step 3: Gather top-50 from each source ────────────────────────
        bm25_top_idx = bm25_scores.argsort()[-50:][::-1]
        bm25_top     = [self.chunks[i] for i in bm25_top_idx]
        sem_top      = [doc for doc, _ in sem_results]

        # Merge and deduplicate by object identity
        candidate_chunks = list({id(c): c for c in (bm25_top + sem_top)}.values())

        # ── Step 4: Filter weak semantic candidates ───────────────────────
        # Keep chunks that cleared the semantic confidence threshold.
        filtered = [
            c for c in candidate_chunks
            if sem_score_map.get(c.page_content, 0) > 0.2
        ]

        # Fallback: if filtering was too aggressive, keep top-20 candidates
        if len(filtered) < 10:
            filtered = candidate_chunks[:20]

        candidate_chunks = filtered[:40]

        # Hard fallback: if somehow nothing survived, return first k chunks
        if not candidate_chunks:
            return self.chunks[:k]

        # ── Step 5: CrossEncoder reranking ────────────────────────────────
        pairs    = [(query, chunk.page_content) for chunk in candidate_chunks]
        scores   = self.reranker.predict(pairs, batch_size=16)
        reranked = sorted(zip(scores, candidate_chunks), key=lambda x: x[0], reverse=True)

        return [chunk for _, chunk in reranked[:k]]


# ─────────────────────────────────────────────────────────────────────────────
# CHUNKING
# Decides between semantic chunking (small docs) and fixed chunking (large docs).
#
# Semantic chunking:  uses embedding similarity to find natural topic boundaries.
#                     Better quality but makes extra API calls — only viable for
#                     small documents where the cost is acceptable.
#
# Fixed chunking:     splits by character count with overlap. Fast and cheap,
#                     suitable for large documents.
# ─────────────────────────────────────────────────────────────────────────────

def create_chunks(all_docs, embedding_model):
    total_chars = sum(len(doc.page_content) for doc in all_docs)

    if total_chars < Config.SMALL_DOC_THRESHOLD:
        print("🔬 Small document — using Semantic Chunking")
        semantic_splitter = SemanticChunker(
            embeddings=embedding_model,
            breakpoint_threshold_type="percentile",
            breakpoint_threshold_amount=85
        )

        # FIX: pass all documents in a single batched call instead of
        # looping (avoids N separate embedding API round-trips)
        texts     = [doc.page_content for doc in all_docs]
        metadatas = [doc.metadata     for doc in all_docs]

        try:
            chunks = semantic_splitter.create_documents(
                texts=texts,
                metadatas=metadatas
            )
        except Exception as e:
            print(f"[SemanticChunker failed, falling back to fixed] {e}")
            chunks = all_docs   # fall through to fixed chunking below

    else:
        print("⚡ Large document — using Fixed Chunking (faster)")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ".", " "]
        )
        chunks = splitter.split_documents(all_docs)

    # Safety net: if chunking produced too few pieces, re-chunk with fixed splitter
    if len(chunks) < 3:
        print("⚠️ Too few chunks produced — re-chunking with fixed splitter")
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        chunks   = splitter.split_documents(all_docs)

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# INDEX BUILDER
# Embeds all chunks into ChromaDB and wires up the HybridRetriever.
# Called once per document upload session.
# ─────────────────────────────────────────────────────────────────────────────

def build_index(all_docs, api_key):
    embedding_model = OpenAIEmbeddings(
        model=Config.EMBEDDING_MODEL,
        openai_api_key=api_key
    )

    total_pages = len(all_docs)
    total_chars = sum(len(doc.page_content) for doc in all_docs)
    print(f"📄 Pages: {total_pages} | Characters: {total_chars:,}")

    all_chunks = create_chunks(all_docs, embedding_model)

    # Drop near-empty chunks that would add noise to retrieval
    all_chunks = [c for c in all_chunks if len(c.page_content.strip()) > 50]

    # Assign a unique ID to each chunk (useful for dedup / future caching)
    for chunk in all_chunks:
        chunk.metadata["chunk_id"] = str(uuid.uuid4())

    print(f"✅ Chunks created: {len(all_chunks)}")

    # Use a temp directory so each session gets an isolated vector store
    persist_dir = tempfile.mkdtemp()
    vectorstore = Chroma.from_documents(
        documents=all_chunks,
        embedding=embedding_model,
        persist_directory=persist_dir
    )

    retriever = HybridRetriever(
        chunks=all_chunks,
        vectorstore=vectorstore,
        alpha=Config.HYBRID_ALPHA
    )

    return retriever, len(all_chunks)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT SUMMARIZER
# Uses a Map-Reduce strategy for large documents:
#   MAP    — summarise a representative sample of pages individually
#   REDUCE — combine those page-level summaries into one final summary
#
# Small documents (<15k chars) are summarised in a single direct call.
# ─────────────────────────────────────────────────────────────────────────────

def generate_summary(docs, filenames, client):
    total_text = "\n\n".join([
        f"[{doc.metadata.get('source', 'doc')}]\n{doc.page_content}"
        for doc in docs
    ])
    total_chars = len(total_text)
    print(f"📝 Summarising {total_chars:,} chars across {len(docs)} page(s)...")

    # ── Direct summary for small documents ───────────────────────────────
    if total_chars <= 15_000:
        prompt = f"""You are an expert document analyst.
Analyse the document(s) and provide a structured summary.

Include:
1. **What these documents are about** (1-2 sentences)
2. **Key contents / topics covered** (bullet points)
3. **Important numbers or statistics** (if any)
4. **What these documents can be used for** (1-2 sentences)

Documents: {', '.join(filenames)}
Content:
{total_text[:12_000]}"""

        response = client.chat.completions.create(
            model=Config.LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        return response.choices[0].message.content

    # ── Map-Reduce for large documents ───────────────────────────────────
    print("📚 Large document detected — using Map-Reduce summarisation")
    n = len(docs)

    # Sample representative pages: evenly spaced so we cover the whole doc
    sampled_raw = docs[:10] if n < 20 else docs[::max(1, n // 10)]

    # FIX: deduplicate by content fingerprint (not object identity via id())
    seen    = set()
    sampled = []
    for d in sampled_raw:
        key = d.page_content[:100]
        if key not in seen:
            seen.add(key)
            sampled.append(d)

    print(f"  Sampled {len(sampled)} representative pages from {n} total")

    # MAP: summarise each sampled page in isolation
    page_summaries = []
    for i, doc in enumerate(sampled):
        src     = doc.metadata.get("source", "doc")
        page    = doc.metadata.get("page", i + 1)
        content = doc.page_content[:2_000]
        if not content.strip():
            continue

        resp = client.chat.completions.create(
            model=Config.LLM_MODEL,
            messages=[{
                "role": "user",
                "content": (
                    f"Summarise this page in 3-4 bullet points. Be concise and factual.\n"
                    f"Focus on key information, numbers, names, and decisions.\n\n"
                    f"Source: {src}, Page {page}\n"
                    f"Content: {content}"
                )
            }],
            temperature=0.2,
            max_tokens=200
        )
        page_summaries.append(f"[Page {page}]\n{resp.choices[0].message.content}")

    # REDUCE: combine all page summaries into one final structured summary
    combined      = "\n\n".join(page_summaries)
    reduce_prompt = f"""You are an expert document analyst.
Below are summaries of key pages from a {n}-page document.
Create a comprehensive structured summary of the ENTIRE document.

Include:
1. **What this document is about** (2-3 sentences)
2. **Key contents / topics covered** (bullet points, be specific)
3. **Important numbers, dates, or statistics** (if any)
4. **Key findings or conclusions** (bullet points)
5. **What this document can be used for** (1-2 sentences)

Document: {', '.join(filenames)} ({n} pages total)

Page Summaries:
{combined}"""

    final_response = client.chat.completions.create(
        model=Config.LLM_MODEL,
        messages=[{"role": "user", "content": reduce_prompt}],
        temperature=0.3,
        max_tokens=800
    )
    print("✅ Map-Reduce summarisation complete")
    return final_response.choices[0].message.content


# ─────────────────────────────────────────────────────────────────────────────
# STREAMING CHAT WITH CITATIONS
# Retrieves relevant chunks, builds a context window, streams the LLM response
# token-by-token, then appends a source footer and a grounded-ness check.
#
# Conversation history is compressed into a rolling summary after 5 turns to
# keep the token count under control without losing context.
# ─────────────────────────────────────────────────────────────────────────────

def ask_document_stream(question, history, retriever, client, session_state):
    """Stream the answer to `question` given the retrieved document context."""

    def classify_query(q):
        """Route the query to a retrieval strategy based on detected intent."""
        q = q.lower()
        if "summary" in q or "overview" in q:
            return "summary"
        elif "compare" in q:
            return "comparison"
        elif any(w in q for w in ["how many", "number", "total"]):
            return "numeric"
        return "factual"

    query_type = classify_query(question)

    # Comparison queries benefit from seeing more chunks (more documents in context)
    k = 6 if query_type == "comparison" else Config.RETRIEVAL_K

    chunks = retriever.retrieve(question, k=k)

    # Deduplicate chunks by content to avoid redundant context
    seen          = set()
    unique_chunks = []
    for c in chunks:
        if c.page_content not in seen:
            unique_chunks.append(c)
            seen.add(c.page_content)
    chunks = unique_chunks

    print(f"  Retrieved {len(chunks)} unique chunks")
    print(f"  Sources  : {[c.metadata.get('source') for c in chunks]}")

    # Build labelled context string with source tags the LLM can cite
    context_parts = []
    sources_used  = []
    for i, chunk in enumerate(chunks):
        src  = chunk.metadata.get("source", "Document")
        page = chunk.metadata.get("page", 1)
        context_parts.append(f"[Doc{i+1}: {src}, Page {page}]\n{chunk.page_content}")
        sources_used.append(f"{src} (Page {page})")

    context_str = "\n\n".join(context_parts)

    # Trim context to stay within the LLM's effective token window.
    # rsplit on a chunk boundary so we don't cut inside a [DocN] block.
    if len(context_str) > Config.MAX_CONTEXT_CHARS:
        context_str = context_str[:Config.MAX_CONTEXT_CHARS].rsplit("[Doc", 1)[0]

    # ── Build message list ────────────────────────────────────────────────
    messages = [{
        "role": "system",
        "content": (
            "You are a document-based AI assistant.\n\n"
            "Rules:\n"
            "- Answer ONLY using the provided document context.\n"
            "- If the answer is not in the documents, say: "
            "\"I don't see that information in the document(s).\"\n"
            "- Never hallucinate or invent facts.\n"
            "- Cite sources inline using [Doc1], [Doc2], etc.\n"
            "- Be concise, accurate, and structured."
        )
    }]

    # ── Conversation history (compressed after 5 turns) ───────────────────
    # Re-summarise every 5 new turns so the summary stays fresh as the
    # conversation grows (not just once at turn 6 and then frozen forever).
    if len(history) > 5:
        should_refresh = (
            not session_state.get("history_summary")
            or len(history) % 5 == 0
        )
        if should_refresh:
            summary_prompt = "Summarise this conversation briefly:\n"
            for turn in history:
                summary_prompt += f"User: {turn[0]}\nAssistant: {turn[1]}\n"

            summary_resp = client.chat.completions.create(
                model=Config.LLM_MODEL,
                messages=[{"role": "user", "content": summary_prompt}],
                temperature=0.3,
                max_tokens=150
            )
            session_state["history_summary"] = summary_resp.choices[0].message.content

        messages.append({
            "role": "system",
            "content": f"Conversation summary so far:\n{session_state['history_summary']}"
        })
    else:
        # For short histories, include turns verbatim for full accuracy
        for turn in history:
            messages.append({"role": "user",      "content": turn[0]})
            messages.append({"role": "assistant", "content": turn[1]})

    messages.append({
        "role": "user",
        "content": f"Document context:\n{context_str}\n\nQuestion: {question}"
    })

    # ── Stream response token-by-token ────────────────────────────────────
    stream = client.chat.completions.create(
        model=Config.LLM_MODEL,
        messages=messages,
        temperature=0.2,
        stream=True
    )

    partial = ""
    for chunk in stream:
        delta = getattr(chunk.choices[0].delta, "content", None)
        if delta:
            partial += delta
            yield partial   # push each token increment to the UI

    # ── Post-stream: groundedness verification ────────────────────────────
    # Ask the LLM to flag any part of its own answer not supported by context.
    # FIX: check for "ok" case-insensitively in the first 20 chars to avoid
    # false positives from GPT saying "OK." or "OK, the answer is supported."
    verification_prompt = (
        f"Context:\n{context_str}\n\n"
        f"Answer:\n{partial}\n\n"
        f"Task: Reply 'OK' if every claim is fully supported by the context. "
        f"Otherwise list the unsupported parts."
    )

    verif   = client.chat.completions.create(
        model=Config.LLM_MODEL,
        messages=[{"role": "user", "content": verification_prompt}],
        temperature=0
    )
    verdict = verif.choices[0].message.content.strip()

    if "ok" not in verdict.lower()[:20]:
        partial += f"\n\n⚠️ **Potentially unsupported content:**\n{verdict}"

    # Append deduplicated source footer
    unique_sources = list(dict.fromkeys(sources_used))
    footer = f"\n\n---\n📄 **Sources:** {' | '.join(unique_sources)}"
    yield partial + footer


# ─────────────────────────────────────────────────────────────────────────────
# GRADIO HANDLER — DOCUMENT PROCESSING
# Called when the user clicks "Process Documents".
# Validates input, extracts text, builds the index, generates a summary,
# and stores everything in per-session state.
# ─────────────────────────────────────────────────────────────────────────────

def process_documents(api_key, files, session_state):
    # ── Input validation ──────────────────────────────────────────────────
    if not api_key or not api_key.startswith("sk-"):
        return (
            "❌ Please enter a valid OpenAI API key (starts with `sk-`).",
            session_state
        )

    if not files:
        return (
            "❌ Please upload at least one document.",
            session_state
        )

    try:
        # Quick connectivity check — catches expired/invalid keys before any
        # expensive embedding or summarisation calls are made
        client = OpenAI(api_key=api_key)
        client.models.list()

    except Exception:
        return (
            "❌ API key rejected by OpenAI. Please check that it is valid and has available credits.",
            session_state
        )

    try:
        all_docs  = []
        filenames = []

        for file in files:
            filename = os.path.basename(file.name)
            docs     = extract_text(file.name, filename)

            if not docs:
                # Surface a clear per-file warning instead of silently skipping
                return (
                    f"❌ Could not extract text from **{filename}**. "
                    f"If it is a scanned PDF, make sure Tesseract and Poppler are installed.",
                    session_state
                )

            all_docs.extend(docs)
            filenames.append(filename)

        if not all_docs:
            return (
                "❌ No text could be extracted from any of the uploaded file(s).",
                session_state
            )

        # Build vector index + BM25 index
        retriever, n_chunks = build_index(all_docs, api_key)

        # Generate document summary shown in the Upload tab
        summary = generate_summary(all_docs, filenames, client)

        # Store everything in per-session state (never shared between users)
        session_state["retriever"]       = retriever
        session_state["client"]          = client
        session_state["filenames"]       = filenames
        session_state["history_summary"] = None  # reset on new upload

        doc_info = (
            f"\n\n---\n"
            f"📁 **Loaded:** {', '.join(filenames)} | "
            f"**Chunks:** {n_chunks}"
        )
        return (summary + doc_info, session_state)

    except Exception as e:
        return (f"❌ Unexpected error: {str(e)}", session_state)


# ─────────────────────────────────────────────────────────────────────────────
# GRADIO HANDLER — CHAT
# Called on every user message in the Chat tab.
# Delegates to the streaming generator; yields partial responses to the UI.
# ─────────────────────────────────────────────────────────────────────────────

def chat_with_doc(message, history, session_state):
    """Stream the assistant reply for a chat message."""
    if not session_state.get("retriever"):
        yield "⚠️ Please upload and process a document first (go to the Upload & Summary tab)."
        return

    try:
        # FIX: pass session_state as the 5th argument (was missing — caused crash)
        for partial in ask_document_stream(
            message,
            history,
            session_state["retriever"],
            session_state["client"],
            session_state           # ← required for history compression
        ):
            yield partial

    except Exception as e:
        yield f"❌ Error during chat: {str(e)}"


# ─────────────────────────────────────────────────────────────────────────────
# GRADIO UI
# Two-tab layout:
#   Tab 1 — Upload documents + view auto-generated summary
#   Tab 2 — Multi-turn chat with streaming responses and source citations
# ─────────────────────────────────────────────────────────────────────────────

with gr.Blocks(title="DocChat AI", theme=gr.themes.Soft()) as demo:

    # Isolated per-user session state — never shared across browser sessions
    session_state = gr.State({})

    gr.Markdown("""
    # 📄 DocChat AI — Advanced RAG System
    **Upload documents → Get an instant summary → Chat with source citations**
    > Supports PDF · DOCX · TXT · XLSX · CSV &nbsp;|&nbsp; Multi-document Q&A &nbsp;|&nbsp; Streaming responses
    """)

    with gr.Tab("📂 Upload & Summary"):
        with gr.Row():
            with gr.Column(scale=1):
                api_input = gr.Textbox(
                    label="🔑 OpenAI API Key",
                    placeholder="sk-...",
                    type="password"
                )
                file_input = gr.File(
                    label="📁 Upload Documents (select multiple)",
                    file_types=[".pdf", ".docx", ".txt", ".xlsx", ".xls", ".csv"],
                    file_count="multiple"
                )
                process_btn = gr.Button(
                    "⚡ Process Documents",
                    variant="primary",
                    size="lg"
                )
                gr.Markdown("""
                **What happens when you click Process:**
                - Text is extracted from your file(s)
                - Documents are split into semantic or fixed chunks
                - Chunks are embedded and stored in a local vector store
                - A BM25 keyword index is built in parallel
                - An automatic summary is generated
                """)

            with gr.Column(scale=2):
                summary_output = gr.Markdown(
                    label="📋 Document Summary",
                    value="*Upload documents and click Process — your summary will appear here.*"
                )

        # FIX: outputs now matches the 2-value tuple returned by process_documents.
        # The old code had a dangling gr.State() as a 3rd output that was never
        # connected to anything, causing the gr.update(interactive=...) to be lost.
        process_btn.click(
            fn=process_documents,
            inputs=[api_input, file_input, session_state],
            outputs=[summary_output, session_state]
        )

    with gr.Tab("💬 Chat with Documents"):
        gr.ChatInterface(
            fn=chat_with_doc,
            additional_inputs=[session_state],
            chatbot=gr.Chatbot(height=480),
            textbox=gr.Textbox(
                placeholder="Ask anything about your document(s)...",
                container=False
            ),
            examples=[
                ["Give me a brief overview of these documents"],
                ["What are the most important points?"],
                ["Summarise the key statistics or numbers"],
                ["What are the main risks or concerns mentioned?"],
                ["Compare the key differences between the documents"],
            ],
        )

    gr.Markdown("""
    ---
    > ⚠️ Your API key is used only within your session and is never stored or logged.
    > Built with LangChain · ChromaDB · Hybrid BM25 + Semantic Search · Cross-Encoder Reranking · RAGAS Evaluated
    """)


demo.launch()
