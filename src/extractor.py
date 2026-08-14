# TEXT EXTRACTION
# One entry point per supported format. PDFs get a fallback chain because no
# single parser handles every PDF in the wild: PyPDFLoader covers the common
# case, PyMuPDF recovers unusual encodings and embedded fonts, pdfplumber
# handles table-heavy layouts, and OCR is the last resort for scans.

import os

import docx
import openpyxl
import pandas as pd
from langchain_core.documents import Document

from src.logging_setup import get_logger

log = get_logger(__name__)

SUPPORTED_EXTENSIONS = ("pdf", "docx", "txt", "csv", "xlsx", "xls")

# Rows sampled from a CSV when building its text representation. Enough to
# convey shape and typical values without blowing up the context window.
CSV_SAMPLE_ROWS = 100


class UnsupportedFormatError(ValueError):
    """Raised when a file extension has no registered extractor."""


def extract_text(filepath, filename):
    """Extract `filename` into a list of Documents, one per logical page.

    Returns an empty list when the file contains no readable text — callers
    are expected to surface that to the user rather than treat it as success.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    handlers = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "txt": _extract_txt,
        "xlsx": _extract_excel,
        "xls": _extract_excel,
        "csv": _extract_csv,
    }

    handler = handlers.get(ext)
    if handler is None:
        raise UnsupportedFormatError(
            f"Unsupported file type '.{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    docs = handler(filepath, filename)

    if not docs:
        log.warning("No text extracted from %s", filename)
    else:
        log.info("Extracted %d page(s) from %s", len(docs), filename)

    return docs


def _make_doc(text, filename, page):
    return Document(
        page_content=text,
        metadata={"source": filename, "page": page, "doc_name": filename},
    )


# ── PDF ───────────────────────────────────────────────────────────────────

def _extract_pdf(filepath, filename):
    """Try each PDF parser in order of speed, returning the first that works."""
    for attempt in (_pdf_pypdf, _pdf_pymupdf, _pdf_plumber, _pdf_ocr):
        try:
            docs = attempt(filepath, filename)
        except Exception as exc:
            log.warning("%s failed on %s: %s", attempt.__name__, filename, exc)
            continue

        if docs:
            log.info("Parsed %s using %s", filename, attempt.__name__)
            return docs

    return []


def _pdf_pypdf(filepath, filename):
    from langchain_community.document_loaders import PyPDFLoader

    pages = PyPDFLoader(filepath).load()
    pages = [p for p in pages if p.page_content and p.page_content.strip()]

    for i, page in enumerate(pages):
        page.metadata["source"] = filename
        page.metadata["doc_name"] = filename
        page.metadata.setdefault("page", i + 1)

    return pages


def _pdf_pymupdf(filepath, filename):
    import fitz  # PyMuPDF

    docs = []
    with fitz.open(filepath) as pdf:
        for i, page in enumerate(pdf):
            text = page.get_text()
            if text and text.strip():
                docs.append(_make_doc(text, filename, i + 1))

    return docs


def _pdf_plumber(filepath, filename):
    """Layout-aware extraction — recovers text from table-heavy PDFs."""
    import pdfplumber

    docs = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(_make_doc(text, filename, i + 1))

    return docs


def _pdf_ocr(filepath, filename):
    """Last resort for scanned PDFs. Requires Tesseract + Poppler on PATH."""
    import pytesseract
    from pdf2image import convert_from_path

    docs = []
    for i, image in enumerate(convert_from_path(filepath, dpi=200)):
        text = pytesseract.image_to_string(image)
        if text and text.strip():
            docs.append(_make_doc(text, filename, i + 1))

    return docs


# ── Other formats ─────────────────────────────────────────────────────────

def _extract_docx(filepath, filename):
    document = docx.Document(filepath)

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables hold a lot of the substance in reports and contracts, and are
    # invisible to paragraph iteration.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    return [_make_doc(text, filename, 1)] if text.strip() else []


def _extract_txt(filepath, filename):
    with open(filepath, encoding="utf-8", errors="ignore") as handle:
        text = handle.read()

    return [_make_doc(text, filename, 1)] if text.strip() else []


def _extract_excel(filepath, filename):
    workbook = openpyxl.load_workbook(filepath, data_only=True)

    docs = []
    for page, sheet_name in enumerate(workbook.sheetnames, start=1):
        sheet = workbook[sheet_name]

        lines = [f"[Sheet: {sheet_name}]"]
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None)
            if row_text.strip():
                lines.append(row_text)

        # One Document per sheet, so citations can name the sheet a fact
        # came from instead of pointing at the whole workbook.
        if len(lines) > 1:
            docs.append(_make_doc("\n".join(lines), filename, page))

    return docs


def _extract_csv(filepath, filename):
    try:
        frame = pd.read_csv(filepath, on_bad_lines="skip")
    except Exception as exc:
        log.warning("Could not read CSV %s: %s", filename, exc)
        return []

    if frame.empty:
        return []

    sample = frame.sample(min(len(frame), CSV_SAMPLE_ROWS), random_state=42)

    text = (
        f"Dataset: {filename}\n"
        f"Rows: {len(frame)}\n"
        f"Columns: {', '.join(map(str, frame.columns.tolist()))}\n\n"
        f"{sample.to_string()}"
    )

    return [_make_doc(text, filename, 1)]


def validate_upload(filepath, filename, max_mb):
    """Raise if an upload is unsupported or too large to process."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file type '.{ext}'. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    size_mb = os.path.getsize(filepath) / (1024 * 1024)
    if size_mb > max_mb:
        raise ValueError(
            f"{filename} is {size_mb:.1f} MB, over the {max_mb} MB limit."
        )
